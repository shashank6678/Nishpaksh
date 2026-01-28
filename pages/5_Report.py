# pages/5_Report.py
# FINAL REPORT COMPILER — WIRE-FRAME ALIGNED (TEC 7.1 COMPLIANT)
# UI REFINED + SECTION STATUS (NO LOGIC CHANGES)

import streamlit as st
import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Inches
import tempfile
import matplotlib.pyplot as plt

# ==================================================
# PAGE CONFIG
# ==================================================
st.set_page_config(layout="wide")
st.title("Final Fairness Evaluation Report")
st.markdown(
    """
    <style>
    /* ======================================================
       IMPORTS & ROOT VARIABLES
       ====================================================== */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');

    :root {
        --primary: #2563eb;
        --primary-dark: #1e40af;
        --primary-light: #3b82f6;
        --accent: #0ea5e9;
        --success: #10b981;
        --warning: #f59e0b;
        --danger: #ef4444;
        
        --bg-app: #f8fafc;
        --bg-card: #ffffff;
        --bg-section: #f1f5f9;
        --bg-hover: #e0f2fe;
        
        --border: #e2e8f0;
        --border-strong: #cbd5e1;
        --shadow-sm: 0 1px 3px 0 rgba(0, 0, 0, 0.1);
        --shadow-md: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
        --shadow-lg: 0 10px 15px -3px rgba(0, 0, 0, 0.1);
        
        --text-primary: #0f172a;
        --text-secondary: #475569;
        --text-muted: #64748b;
        --text-light: #94a3b8;
    }

    /* ======================================================
       GLOBAL BASE
       ====================================================== */
    * {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Helvetica, Arial, sans-serif !important;
    }

    .stApp {
        background: linear-gradient(135deg, #f8fafc 0%, #e0f2fe 100%);
        color: var(--text-primary);
    }

    html, body {
        font-size: 16px;
        line-height: 1.6;
        -webkit-font-smoothing: antialiased;
        -moz-osx-font-smoothing: grayscale;
    }

    /* ======================================================
       TYPOGRAPHY
       ====================================================== */
    h1 {
        font-size: 2.5rem !important;
        font-weight: 700 !important;
        letter-spacing: -0.025em !important;
        margin-bottom: 0.5rem !important;
        color: var(--text-primary) !important;
        background: linear-gradient(135deg, #2563eb 0%, #0ea5e9 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
    }

    h1::after {
        content: "";
        display: block;
        width: 80px;
        height: 4px;
        background: linear-gradient(90deg, var(--primary) 0%, var(--accent) 100%);
        margin-top: 12px;
        border-radius: 2px;
        box-shadow: 0 2px 4px rgba(37, 99, 235, 0.3);
    }

    h2 {
        font-size: 1.75rem !important;
        font-weight: 700 !important;
        margin-top: 2.5rem !important;
        margin-bottom: 1rem !important;
        color: var(--text-primary) !important;
        padding-left: 1rem !important;
        border-left: 5px solid var(--primary) !important;
        background: linear-gradient(90deg, rgba(37, 99, 235, 0.05) 0%, transparent 100%);
        padding: 0.75rem 0 0.75rem 1rem !important;
        border-radius: 0 8px 8px 0;
    }

    h3 {
        font-size: 1.25rem !important;
        font-weight: 600 !important;
        color: var(--text-primary) !important;
        margin-bottom: 0.75rem !important;
    }

    p, .stMarkdown p {
        color: var(--text-secondary) !important;
        font-size: 1rem !important;
        line-height: 1.7 !important;
        max-width: 75ch;
    }

    .stCaption {
        color: var(--text-muted) !important;
        font-size: 0.9rem !important;
    }

    /* ======================================================
       SURVEY NAVIGATION SIDEBAR
       ====================================================== */
    div[data-testid="stVerticalBlockBorderWrapper"]:has(button) {
        background: var(--bg-card) !important;
        border: 2px solid var(--border) !important;
        border-radius: 12px !important;
        padding: 1.5rem !important;
        margin-bottom: 1.5rem !important;
        box-shadow: var(--shadow-md) !important;
    }

    div[data-testid="stVerticalBlockBorderWrapper"] h3 {
        font-size: 1.1rem !important;
        font-weight: 600 !important;
        color: var(--text-primary) !important;
        margin-bottom: 1rem !important;
        padding-bottom: 0.75rem !important;
        border-bottom: 2px solid var(--border) !important;
    }

    /* Navigation buttons */
    div[data-testid="stVerticalBlockBorderWrapper"] button {
        width: 100% !important;
        text-align: left !important;
        padding: 0.75rem 1rem !important;
        margin-bottom: 0.5rem !important;
        border: 1px solid var(--border) !important;
        border-radius: 8px !important;
        background: var(--bg-section) !important;
        color: var(--text-secondary) !important;
        font-size: 0.95rem !important;
        font-weight: 500 !important;
        transition: all 0.2s ease !important;
    }

    div[data-testid="stVerticalBlockBorderWrapper"] button:hover {
        background: var(--bg-hover) !important;
        border-color: var(--primary) !important;
        color: var(--primary) !important;
        transform: translateX(4px);
        box-shadow: var(--shadow-sm) !important;
    }

    /* ======================================================
       MAIN CONTENT CARDS
       ====================================================== */
    div[data-testid="stVerticalBlock"] > div[data-testid="stVerticalBlockBorderWrapper"] {
        background: var(--bg-card) !important;
        border: 1px solid var(--border) !important;
        border-radius: 16px !important;
        padding: 2rem !important;
        margin-bottom: 1.5rem !important;
        box-shadow: var(--shadow-md) !important;
        transition: all 0.3s ease;
    }

    div[data-testid="stVerticalBlock"] > div[data-testid="stVerticalBlockBorderWrapper"]:hover {
        box-shadow: var(--shadow-lg) !important;
        border-color: var(--border-strong) !important;
    }

    /* ======================================================
       SURVEY QUESTIONS (EXPANDERS) - KEY IMPROVEMENTS
       ====================================================== */
    div[data-testid="stExpander"] {
        margin-bottom: 1.25rem !important;
        border-radius: 12px !important;
        overflow: hidden;
        box-shadow: var(--shadow-sm) !important;
        transition: all 0.3s ease;
    }

    div[data-testid="stExpander"]:hover {
        box-shadow: var(--shadow-md) !important;
    }

    /* Question headers */
    div[data-testid="stExpander"] > details > summary {
        font-size: 1.05rem !important;
        font-weight: 600 !important;
        color: var(--text-primary) !important;
        background: linear-gradient(135deg, #f8fafc 0%, #f1f5f9 100%) !important;
        padding: 1.25rem 1.5rem !important;
        border-radius: 12px !important;
        border: 2px solid var(--border) !important;
        border-left: 6px solid var(--primary) !important;
        cursor: pointer !important;
        transition: all 0.2s ease !important;
        line-height: 1.6 !important;
        min-height: 60px;
        display: flex;
        align-items: center;
    }

    div[data-testid="stExpander"] > details > summary:hover {
        background: linear-gradient(135deg, var(--bg-hover) 0%, #e0f2fe 100%) !important;
        border-left-color: var(--accent) !important;
        transform: translateX(4px);
        box-shadow: var(--shadow-sm) !important;
    }

    div[data-testid="stExpander"] > details[open] > summary {
        background: linear-gradient(135deg, var(--primary) 0%, var(--primary-light) 100%) !important;
        color: white !important;
        border-color: var(--primary-dark) !important;
        border-left-color: var(--accent) !important;
        border-radius: 12px 12px 0 0 !important;
    }

    /* Question content area */
    div[data-testid="stExpander"] > details > div {
        background-color: var(--bg-card) !important;
        padding: 1.5rem 1.75rem !important;
        border-left: 6px solid var(--border-strong) !important;
        border-right: 2px solid var(--border) !important;
        border-bottom: 2px solid var(--border) !important;
        border-radius: 0 0 12px 12px !important;
        box-shadow: inset 0 2px 4px rgba(0, 0, 0, 0.02);
    }

    /* ======================================================
       FORM INPUTS & CONTROLS
       ====================================================== */
    label {
        font-weight: 600 !important;
        color: var(--text-primary) !important;
        font-size: 0.95rem !important;
        margin-bottom: 0.5rem !important;
        display: block !important;
    }

    /* Text inputs and textareas */
    textarea, 
    input[type="text"],
    input[type="number"] {
        font-size: 1rem !important;
        line-height: 1.6 !important;
        border-radius: 8px !important;
        padding: 0.75rem 1rem !important;
        border: 2px solid var(--border) !important;
        background: var(--bg-card) !important;
        color: var(--text-primary) !important;
        transition: all 0.2s ease !important;
        width: 100% !important;
    }

    textarea:focus,
    input[type="text"]:focus,
    input[type="number"]:focus {
        border-color: var(--primary) !important;
        box-shadow: 0 0 0 3px rgba(37, 99, 235, 0.1) !important;
        outline: none !important;
    }

    /* Radio buttons and checkboxes */
    div[data-testid="stRadio"] > div,
    div[data-testid="stCheckbox"] > div {
        background: var(--bg-section) !important;
        padding: 1rem !important;
        border-radius: 8px !important;
        border: 1px solid var(--border) !important;
        margin-bottom: 0.5rem !important;
        transition: all 0.2s ease !important;
    }

    div[data-testid="stRadio"] > div:hover,
    div[data-testid="stCheckbox"] > div:hover {
        background: var(--bg-hover) !important;
        border-color: var(--primary) !important;
    }

    /* Radio button labels */
    div[data-testid="stRadio"] label {
        font-weight: 500 !important;
        font-size: 0.95rem !important;
        padding: 0.5rem 0 !important;
    }

    /* ======================================================
       BUTTONS
       ====================================================== */
    button[kind="primary"] {
        background: linear-gradient(135deg, var(--primary) 0%, var(--primary-light) 100%) !important;
        color: white !important;
        font-weight: 600 !important;
        font-size: 1rem !important;
        border-radius: 10px !important;
        padding: 0.875rem 2rem !important;
        border: none !important;
        box-shadow: var(--shadow-md) !important;
        transition: all 0.3s ease !important;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }

    button[kind="primary"]:hover {
        background: linear-gradient(135deg, var(--primary-dark) 0%, var(--primary) 100%) !important;
        transform: translateY(-2px);
        box-shadow: var(--shadow-lg) !important;
    }

    button[kind="secondary"] {
        background: var(--bg-card) !important;
        color: var(--text-primary) !important;
        border: 2px solid var(--border) !important;
        font-weight: 600 !important;
        border-radius: 10px !important;
        padding: 0.875rem 2rem !important;
        transition: all 0.2s ease !important;
    }

    button[kind="secondary"]:hover {
        border-color: var(--primary) !important;
        color: var(--primary) !important;
        background: var(--bg-hover) !important;
    }

    /* ======================================================
       METRICS & INFO BOXES
       ====================================================== */
    div[data-testid="stMetric"] {
        background: linear-gradient(135deg, var(--bg-card) 0%, var(--bg-section) 100%) !important;
        padding: 1.5rem !important;
        border-radius: 12px !important;
        border: 2px solid var(--border) !important;
        box-shadow: var(--shadow-md) !important;
    }

    div[data-testid="stMetric"] label {
        font-size: 0.9rem !important;
        font-weight: 600 !important;
        color: var(--text-secondary) !important;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }

    div[data-testid="stMetric"] div[data-testid="stMetricValue"] {
        font-size: 2.5rem !important;
        font-weight: 700 !important;
        color: var(--primary) !important;
    }

    /* Info/Success/Warning boxes */
    div[data-testid="stAlert"] {
        border-radius: 12px !important;
        border: none !important;
        box-shadow: var(--shadow-sm) !important;
        padding: 1.25rem 1.5rem !important;
    }

    div[data-testid="stAlert"][data-baseweb="notification"][kind="info"] {
        background: linear-gradient(135deg, #dbeafe 0%, #e0f2fe 100%) !important;
        border-left: 5px solid var(--accent) !important;
    }

    div[data-testid="stAlert"][data-baseweb="notification"][kind="success"] {
        background: linear-gradient(135deg, #d1fae5 0%, #a7f3d0 100%) !important;
        border-left: 5px solid var(--success) !important;
    }

    div[data-testid="stAlert"][data-baseweb="notification"][kind="warning"] {
        background: linear-gradient(135deg, #fef3c7 0%, #fde68a 100%) !important;
        border-left: 5px solid var(--warning) !important;
    }

    div[data-testid="stAlert"][data-baseweb="notification"][kind="error"] {
        background: linear-gradient(135deg, #fee2e2 0%, #fecaca 100%) !important;
        border-left: 5px solid var(--danger) !important;
    }

    /* ======================================================
       SIDEBAR
       ====================================================== */
    section[data-testid="stSidebar"] {
        background: linear-gradient(180deg, #ffffff 0%, #f8fafc 100%) !important;
        border-right: 2px solid var(--border) !important;
        box-shadow: 2px 0 8px rgba(0, 0, 0, 0.05);
    }

    section[data-testid="stSidebar"] h1,
    section[data-testid="stSidebar"] h2,
    section[data-testid="stSidebar"] h3 {
        color: var(--text-primary) !important;
        padding-left: 0.5rem !important;
    }

    section[data-testid="stSidebar"] li {
        border-radius: 8px !important;
        margin-bottom: 0.25rem !important;
        transition: all 0.2s ease !important;
    }

    section[data-testid="stSidebar"] li:has(a[aria-current="page"]) {
        background: linear-gradient(135deg, var(--primary) 0%, var(--primary-light) 100%) !important;
        box-shadow: var(--shadow-sm) !important;
    }

    section[data-testid="stSidebar"] li:has(a[aria-current="page"]) a {
        color: white !important;
        font-weight: 600 !important;
    }

    section[data-testid="stSidebar"] li:hover {
        background: var(--bg-hover) !important;
    }

    /* ======================================================
       PROGRESS & INDICATORS
       ====================================================== */
    .stProgress > div > div {
        background: linear-gradient(90deg, var(--primary) 0%, var(--accent) 100%) !important;
        border-radius: 10px !important;
        height: 8px !important;
    }

    /* ======================================================
       DIVIDERS
       ====================================================== */
    hr {
        border: none !important;
        border-top: 2px solid var(--border) !important;
        margin: 2.5rem 0 !important;
        opacity: 0.6;
    }

    /* ======================================================
       SCROLLBAR STYLING
       ====================================================== */
    ::-webkit-scrollbar {
        width: 10px;
        height: 10px;
    }

    ::-webkit-scrollbar-track {
        background: var(--bg-section);
        border-radius: 10px;
    }

    ::-webkit-scrollbar-thumb {
        background: linear-gradient(180deg, var(--primary) 0%, var(--accent) 100%);
        border-radius: 10px;
        border: 2px solid var(--bg-section);
    }

    ::-webkit-scrollbar-thumb:hover {
        background: linear-gradient(180deg, var(--primary-dark) 0%, var(--primary) 100%);
    }

    /* ======================================================
       RESPONSIVE IMPROVEMENTS
       ====================================================== */
    @media (max-width: 768px) {
        h1 {
            font-size: 2rem !important;
        }
        
        h2 {
            font-size: 1.5rem !important;
        }
        
        div[data-testid="stExpander"] > details > summary {
            font-size: 0.95rem !important;
            padding: 1rem 1.25rem !important;
        }
    }

    /* ======================================================
       ANIMATIONS
       ====================================================== */
    @keyframes slideIn {
        from {
            opacity: 0;
            transform: translateY(10px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }

    div[data-testid="stVerticalBlock"] > div {
        animation: slideIn 0.3s ease-out;
    }

    /* ======================================================
       FOCUS STATES
       ====================================================== */
    button:focus-visible,
    input:focus-visible,
    textarea:focus-visible {
        outline: 3px solid rgba(37, 99, 235, 0.5) !important;
        outline-offset: 2px !important;
    }
    </style>
    """,
    unsafe_allow_html=True
)



# ==================================================
# PRECONDITIONS
# ==================================================
REQUIRED_KEYS = ["survey_outputs", "preproc", "inference", "results"]
for key in REQUIRED_KEYS:
    if key not in st.session_state:
        st.error(f"Missing required step: {key}")
        st.stop()

survey = st.session_state["survey_outputs"]
preproc = st.session_state["preproc"]
inference = st.session_state["inference"]
results = st.session_state["results"]

# ==================================================
# DEFENSIVE INITIALIZATION (UNCHANGED)
# ==================================================
survey.setdefault("risk_bucket", "Medium")
survey.setdefault("auditor_access", "Training + Validation")
survey.setdefault("testing_type", "Grey box")
survey.setdefault("dependencies", "")
survey.setdefault("limitations", "")
survey.setdefault("questionnaire_summary", "")
survey.setdefault("protected_attr", "")
survey.setdefault("privileged_groups", "")
survey.setdefault("favourable_outcome", "")
survey.setdefault("protected_attr_rationale", "")
survey.setdefault("metric_rationale", "")
survey.setdefault("threshold_rationale", "")
survey.setdefault("risk_outcome", "")
survey.setdefault("certification_context", "")

preproc.setdefault("user_narratives", {})
preproc["user_narratives"].setdefault("P14", "")
preproc.setdefault("pipeline_description", "")
preproc.setdefault("split_method", "Not specified")
preproc.setdefault("synthetic_data", "No synthetic data used.")
preproc.setdefault("scenarios_tested", "")

# ==================================================
# TEMPLATE
# ==================================================
TEMPLATE_PATH = Path(__file__).parent / "Fairness_Evaluation_Report_Wireframe.docx"
if not TEMPLATE_PATH.exists():
    st.error("Wireframe DOCX template not found.")
    st.stop()

# ==================================================
# DOCX HELPERS (UNCHANGED)
# ==================================================
def replace_text(doc, token, value):
    found = False
    for p in doc.paragraphs:
        if token in p.text:
            p.text = p.text.replace(token, str(value))
            found = True
    return found

def insert_table(doc, token, df):
    for p in doc.paragraphs:
        if token in p.text:
            p.text = ""
            table = doc.add_table(rows=1, cols=len(df.columns))
            for i, col in enumerate(df.columns):
                table.rows[0].cells[i].text = str(col)
            for _, row in df.iterrows():
                cells = table.add_row().cells
                for i, v in enumerate(row):
                    cells[i].text = f"{v:.4f}" if isinstance(v, float) else str(v)
            p._p.addnext(table._tbl)
            return True
    return False

def insert_plot(doc, token, fig):
    for p in doc.paragraphs:
        if token in p.text:
            p.text = p.text.replace(token, "")
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                fig.savefig(tmp.name, dpi=240, bbox_inches="tight")
                plt.close(fig)
                p.add_run().add_picture(tmp.name, width=Inches(5.5))
            return True
    return False

# ==================================================
# PLOTS (UNCHANGED)
# ==================================================
def plot_fairness_summary():
    metrics = results["fairness_metrics"]
    fig, ax = plt.subplots(figsize=(6, 3))
    ax.bar(metrics.keys(), metrics.values())
    ax.axhline(0, linestyle="--", color="black")
    ax.axhline(1, linestyle="--", color="black")
    ax.set_title("Bias Index by Protected Attribute")
    plt.xticks(rotation=30, ha="right")
    return fig

def plot_accuracy_bar():
    df = inference["results_df"]
    fig, ax = plt.subplots(figsize=(6, 3))
    ax.bar(df["Model"], df["Accuracy"])
    ax.set_ylim(0, 1)
    ax.set_title("Model Performance Comparison")
    plt.xticks(rotation=30, ha="right")
    return fig

def plot_fairness_ci():
    fb = inference["fairness_bootstrap"]
    fig, ax = plt.subplots(figsize=(6, 3))
    for metric, models in fb.items():
        for model, vals in models.items():
            ax.errorbar(model, np.mean(vals), yerr=np.std(vals), fmt="o")
    ax.set_title("Fairness Metric Uncertainty")
    return fig

# ==================================================
# SECTION STATUS PANEL (RESTORED & EXPLICIT)
# ==================================================
# ==================================================
# SECTION STATUS PANEL (INTENTIONAL COMPLETION CHECK)
# ==================================================
def is_meaningful(text: str, min_len: int = 10) -> bool:
    if text is None:
        return False
    return len(text.strip()) >= min_len

def status_badge(ok: bool) -> str:
    return "🟢 Complete" if ok else "🟠 Needs input"

with st.container(border=True):
    st.markdown("### Report completion status")
    st.caption(
        "Sections are marked complete only after sufficient, explicit information "
        "has been provided. Default selections do not imply completion."
    )

    col1, col2 = st.columns(2)

    with col1:
        st.write(
            f"**1. System context and scope** — "
            f"{status_badge(is_meaningful(survey['protected_attr']) and is_meaningful(survey['favourable_outcome']))}"
        )

        st.write(
            f"**2. Audit configuration** — "
            f"{status_badge(survey['testing_type'] in ['Open box','Grey box','Closed box'])}"
        )

        st.write(
            f"**3. Data and pipeline description** — "
            f"{status_badge(is_meaningful(preproc['pipeline_description']))}"
        )

    with col2:
        st.write(
            f"**4. Risk and limitations** — "
            f"{status_badge(is_meaningful(survey['limitations']))}"
        )

        st.write(
            f"**5. Fairness rationale** — "
            f"{status_badge(is_meaningful(survey['metric_rationale']) and is_meaningful(survey['threshold_rationale']))}"
        )

        st.write(
            f"**6. Certification context** — "
            f"{status_badge(is_meaningful(survey['certification_context']))}"
        )

# ==================================================
# UI TABS
# ==================================================
tab1, tab2, tab3 = st.tabs(["Summary", "Metrics", "Detailed Report"])

# ==================================================
# TAB 1 — SUMMARY
# ==================================================
with tab1:
    st.markdown("### System context and audit scope")
    st.caption(
        "Factual declarations describing what was evaluated and under what conditions."
    )

    with st.container(border=True):
        st.markdown(
            f"""
            Fairness Score: **{results['FS']:.3f}**  
            Bias Index: **{results['BI']:.3f}**  
            Final verdict: **{results['verdict']}**
            """
        )

    with st.container(border=True):
        survey["protected_attr"] = st.text_input(
            "Which protected attributes were evaluated?",
            survey["protected_attr"]
        )

        survey["privileged_groups"] = st.text_input(
            "How were privileged and unprivileged groups defined?",
            survey["privileged_groups"]
        )

        survey["favourable_outcome"] = st.text_input(
            "What outcome was considered favourable?",
            survey["favourable_outcome"]
        )

    with st.container(border=True):
        survey["risk_bucket"] = st.selectbox(
            "Risk classification of this evaluation",
            ["Low", "Medium", "High"],
            index=["Low", "Medium", "High"].index(survey["risk_bucket"])
        )

        survey["auditor_access"] = st.selectbox(
            "Data access available to the evaluator",
            ["Training data only", "Training + Validation", "Full (Train/Validation/Test)"],
            index=1
        )

        survey["testing_type"] = st.radio(
            "System access level during testing",
            ["Open box", "Grey box", "Closed box"],
            index=["Open box", "Grey box", "Closed box"].index(survey["testing_type"])
        )

    with st.container(border=True):
        depends = st.radio(
            "Did this evaluation require assistance from the system developer?",
            ["No", "Yes"],
            index=0 if survey["dependencies"] in ("", "No developer dependency.") else 1
        )

        if depends == "Yes":
            survey["dependencies"] = st.text_area(
                "Describe the nature of developer involvement",
                survey["dependencies"],
                height=120
            )
        else:
            survey["dependencies"] = "No developer dependency."

        survey["limitations"] = st.text_area(
            "Which aspects of the system were not evaluated as part of this assessment?",
            survey["limitations"],
            height=140
        )

# ==================================================
# TAB 2 — METRICS (LOCKED)
# ==================================================
with tab2:
    st.markdown("### Computed metrics")
    st.caption("These values are generated automatically and cannot be edited.")

    with st.container(border=True):
        st.dataframe(inference["tables"]["fairness"], use_container_width=True)

    with st.container(border=True):
        st.dataframe(inference["tables"]["performance"], use_container_width=True)

# ==================================================
# TAB 3 — DETAILED REPORT
# ==================================================
with tab3:
    st.markdown("### Detailed evaluation narrative")

    with st.container(border=True):
        preproc["user_narratives"]["P14"] = st.text_area(
            "Describe the AI system and its intended use",
            preproc["user_narratives"]["P14"],
            height=140
        )

        preproc["pipeline_description"] = st.text_area(
            "Describe the data, model, and processing pipeline",
            preproc["pipeline_description"],
            height=140
        )

    with st.container(border=True):
        survey["questionnaire_summary"] = st.text_area(
            "Summary of fairness questionnaire responses",
            survey["questionnaire_summary"],
            height=120
        )

        survey["risk_outcome"] = st.text_area(
            "Outcome of the risk assessment",
            survey["risk_outcome"],
            height=120
        )

    with st.container(border=True):
        survey["protected_attr_rationale"] = st.text_area(
            "Rationale for selecting protected attributes",
            survey["protected_attr_rationale"],
            height=120
        )

        survey["metric_rationale"] = st.text_area(
            "Rationale for selecting fairness metrics",
            survey["metric_rationale"],
            height=120
        )

        survey["threshold_rationale"] = st.text_area(
            "Rationale for threshold selection",
            survey["threshold_rationale"],
            height=120
        )

    with st.container(border=True):
        preproc["scenarios_tested"] = st.text_area(
            "Evaluation scenarios tested",
            preproc["scenarios_tested"],
            height=120
        )

        survey["certification_context"] = st.text_area(
            "Certification and intended usage context",
            survey["certification_context"],
            height=120
        )

# ==================================================
# FINAL GENERATION (UNCHANGED)
# ==================================================
if st.button("Generate Final Report"):
    doc = Document(TEMPLATE_PATH)

    TEXT = {
        "[[P1_TEXT]]": f"FS={results['FS']:.3f}, BI={results['BI']:.3f}, Verdict={results['verdict']}",
        "[[P2_TEXT]]": preproc["user_narratives"]["P14"],
        "[[P3_TEXT]]": survey["protected_attr"],
        "[[P4_TEXT]]": survey["privileged_groups"],
        "[[P5_TEXT]]": survey["favourable_outcome"],
        "[[P6_TEXT]]": survey["risk_bucket"],
        "[[P7_TEXT]]": survey["auditor_access"],
        "[[P8_TEXT]]": survey["testing_type"],
        "[[P9_TEXT]]": survey["dependencies"],
        "[[P10_TEXT]]": survey["limitations"],
        "[[P13_TEXT]]": f"Overall fairness score FS={results['FS']:.3f} with verdict {results['verdict']}.",
        "[[P14_TEXT]]": preproc["user_narratives"]["P14"],
        "[[P15_TEXT]]": preproc["pipeline_description"],
        "[[P16_TEXT]]": survey["questionnaire_summary"],
        "[[P17_TEXT]]": survey["risk_outcome"],
        "[[P18_TEXT]]": survey["protected_attr_rationale"],
        "[[P19_TEXT]]": survey["metric_rationale"],
        "[[P20_TEXT]]": survey["threshold_rationale"],
        "[[P21_TEXT]]": preproc["split_method"],
        "[[P22_TEXT]]": preproc["synthetic_data"],
        "[[P23_TEXT]]": preproc["scenarios_tested"],
        "[[P26_TEXT]]": survey["certification_context"],
    }

    for k, v in TEXT.items():
        if not replace_text(doc, k, v):
            st.error(f"Missing placeholder in template: {k}")
            st.stop()

    insert_table(doc, "[[TABLE_P11]]", inference["tables"]["fairness"])
    insert_table(doc, "[[TABLE_P24]]", inference["tables"]["performance"])

    insert_plot(doc, "[[FIG_P12]]", plot_fairness_summary())
    insert_plot(doc, "[[FIG_P18]]", plot_accuracy_bar())
    insert_plot(doc, "[[FIG_P19]]", plot_fairness_ci())

    out = Path(tempfile.gettempdir()) / "Fairness_Evaluation_Report_Final.docx"
    doc.save(out)

    st.success("Final report generated successfully.")
    st.download_button("Download Report", open(out, "rb"), file_name=out.name)
