#!/usr/bin/env python3
"""
generate_report.py

Standalone report filler:
- Loads AI_Fairness_Report_Template.docx
- Reads captures/<session> (index.json, inference_metadata.json, results csv, bootstrap npz)
- Replaces text placeholders {{KEY}} with values from metadata or computed summaries
- Replaces {{PLOT:filename.png}} by embedding the PNG from captures (exact or substring match)
- Replaces {{TABLE:NAME}} by inserting a Word table using matching CSV or derived DataFrame
- Saves AI_Fairness_Report_<timestamp>.docx into the session folder and writes report_metadata.json
- Optionally logs to MLflow if installed and MLFLOW_TRACKING_URI is configured

Usage:
    python generate_report.py --session <session_id>
    python generate_report.py --folder /full/path/to/captures/session123
"""

import argparse
import json
import re
import hashlib
from pathlib import Path
from datetime import datetime

import pandas as pd
import numpy as np

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image

# Optional mlflow
try:
    import mlflow
    _MLFLOW_AVAILABLE = True
except Exception:
    _MLFLOW_AVAILABLE = False

PLACEHOLDER_RE = re.compile(r"\{\{([^}]+)\}\}")

def md5_file(path: Path):
    h = hashlib.md5()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()

def find_template(session_folder: Path, template_name="AI_Fairness_Report_Template.docx"):
    # preference order: session captures, cwd, /mnt/data
    cand = session_folder / template_name
    if cand.exists():
        return cand
    cand = Path(template_name)
    if cand.exists():
        return cand
    cand = Path("/mnt/data") / template_name
    if cand.exists():
        return cand
    raise FileNotFoundError(f"Template {template_name} not found in session folder, cwd, or /mnt/data")

def read_index(session_folder: Path):
    idx = session_folder / "index.json"
    if not idx.exists():
        return []
    try:
        return json.loads(idx.read_text(encoding="utf-8"))
    except Exception:
        return []

def find_artifact_by_filename(index_entries, fname_substr):
    # exact match then substring match
    for e in index_entries:
        if e.get("file") == fname_substr or (e.get("abs_path") and Path(e.get("abs_path")).name == fname_substr):
            return Path(e.get("abs_path"))
    for e in index_entries:
        if fname_substr in e.get("file", ""):
            return Path(e.get("abs_path"))
    return None

def insert_image_at_paragraph(paragraph, img_path: Path, max_width_in=6.0):
    try:
        # clear paragraph runs
        for r in list(paragraph.runs):
            r.text = ""
        img = Image.open(img_path)
        dpi = img.info.get("dpi", (300,300))[0] if isinstance(img.info.get("dpi", None), tuple) else (img.info.get("dpi") or 300)
        w_px, h_px = img.size
        width_in = w_px / dpi if dpi and dpi > 0 else max_width_in
        width_in = min(width_in, max_width_in)
        run = paragraph.add_run()
        run.add_picture(str(img_path), width=Inches(width_in))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        img.close()
        return True
    except Exception:
        # leave placeholder text if failed
        return False

def replace_text_placeholders(doc: Document, replacements: dict):
    # paragraphs
    for p in doc.paragraphs:
        if not p.text:
            continue
        def repl(m):
            key = m.group(1)
            return str(replacements.get(key, m.group(0)))
        new_text = PLACEHOLDER_RE.sub(repl, p.text)
        if new_text != p.text:
            try:
                p.clear()  # requires python-docx >=0.8.11
                p.add_run(new_text)
            except Exception:
                # fallback: replace runs
                for run in p.runs:
                    run.text = PLACEHOLDER_RE.sub(repl, run.text)
    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if not cell.text:
                    continue
                def repl2(m):
                    key = m.group(1)
                    return str(replacements.get(key, m.group(0)))
                new_text = PLACEHOLDER_RE.sub(repl2, cell.text)
                if new_text != cell.text:
                    cell.text = new_text

def replace_plot_placeholders(doc: Document, index_entries, session_folder: Path, max_width_in=6.0):
    # find tokens like {{PLOT:filename.png}}
    plot_pat = re.compile(r"\{\{PLOT:([^}]+)\}\}")
    for p in list(doc.paragraphs):
        if not p.text:
            continue
        m = plot_pat.search(p.text)
        if not m:
            continue
        fname = m.group(1).strip()
        img_path = find_artifact_by_filename(index_entries, fname)
        if img_path and img_path.exists():
            inserted = insert_image_at_paragraph(p, img_path, max_width_in=max_width_in)
            if not inserted:
                p.text = p.text.replace(f"{{{{PLOT:{fname}}}}}", f"[Failed to insert image: {fname}]")
        else:
            p.text = p.text.replace(f"{{{{PLOT:{fname}}}}}", f"[Missing plot: {fname}]")

def insert_table_at_placeholder(doc: Document, token_name: str, df: pd.DataFrame, max_rows=50):
    token = f"{{{{TABLE:{token_name}}}}}"
    for p in list(doc.paragraphs):
        if token in p.text:
            if df is None or df.empty:
                p.text = p.text.replace(token, "[No data available]")
                continue
            # create table
            cols = list(df.columns)
            table = doc.add_table(rows=1, cols=len(cols))
            hdr_cells = table.rows[0].cells
            for i, c in enumerate(cols):
                hdr_cells[i].text = str(c)
            for _, row in df.head(max_rows).iterrows():
                cells = table.add_row().cells
                for i, c in enumerate(cols):
                    val = row[c]
                    if isinstance(val, (float, np.floating)):
                        cells[i].text = f"{val:.4f}"
                    else:
                        cells[i].text = str(val)
            # remove paragraph containing token
            try:
                p._element.getparent().remove(p._element)
            except Exception:
                p.text = ""

def load_best_results_csv(session_folder: Path):
    # prefer explicit results_df.csv, else any file containing 'results' and .csv
    candidates = list(session_folder.glob("*results*.csv"))
    exact = session_folder / "results_df.csv"
    if exact.exists():
        try:
            return pd.read_csv(exact)
        except Exception:
            pass
    for c in sorted(candidates, key=lambda p: p.stat().st_mtime, reverse=True):
        try:
            return pd.read_csv(c)
        except Exception:
            continue
    return pd.DataFrame()

def main(session_folder: Path, template_name="AI_Fairness_Report_Template.docx", output_prefix="AI_Fairness_Report"):
    session_folder = session_folder.resolve()
    index_entries = read_index(session_folder)
    # load template
    template_path = find_template(session_folder, template_name=template_name)
    doc = Document(str(template_path))

    # load metadata and results
    inf_meta_path = session_folder / "inference_metadata.json"
    inf_meta = {}
    if inf_meta_path.exists():
        try:
            inf_meta = json.loads(inf_meta_path.read_text(encoding="utf-8"))
        except Exception:
            inf_meta = {}

    results_df = load_best_results_csv(session_folder)
    # optional bootstrap raw arrays (npz)
    bootstrap_path = session_folder / "bootstrap_raw.npz"
    bootstrap = None
    if bootstrap_path.exists():
        try:
            bootstrap = np.load(str(bootstrap_path), allow_pickle=True)
        except Exception:
            bootstrap = None

    # Build replacements from metadata and simple computed fields
    replacements = {}
    # copy keys (uppercased)
    for k, v in inf_meta.items():
        replacements[k.upper()] = v
    # some safe defaults
    replacements.setdefault("RUN_TIMESTAMP", inf_meta.get("timestamp", datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
    replacements.setdefault("SESSION_ID", inf_meta.get("session_id", session_folder.name))
    replacements.setdefault("DATASET_NAME", inf_meta.get("dataset_id", "N/A"))
    replacements.setdefault("TARGET_COL", inf_meta.get("label_col", ""))
    replacements.setdefault("SENSITIVE_COL", inf_meta.get("sensitive_col", ""))
    replacements.setdefault("PRIVILEGED_VALUE", inf_meta.get("privileged_value", ""))
    replacements.setdefault("ARTIFACTS_PATH", str(session_folder))

    # add some computed values
    try:
        df = pd.read_csv(session_folder / "results_df.csv") if (session_folder / "results_df.csv").exists() else results_df
        replacements["DATA_ROWS"] = int(inf_meta.get("n_rows", inf_meta.get("n_samples", ""))) or ""
        replacements["DATA_COLS"] = int(inf_meta.get("n_cols", "")) if inf_meta.get("n_cols") else ""
    except Exception:
        replacements["DATA_ROWS"] = ""
        replacements["DATA_COLS"] = ""

    # simple verdict heuristic (if available)
    if not results_df.empty and "Disparate Impact" in results_df.columns:
        try:
            bad = results_df[results_df["Disparate Impact"].astype(float) < 0.8]
            replacements["VERDICT"] = "REVIEW REQUIRED" if not bad.empty else "PASS"
            replacements["VERDICT_CONFIDENCE"] = ""
            replacements["VERDICT_REASON"] = f"{len(bad)} model(s) have DI<0.8" if not bad.empty else "No DI breaches"
        except Exception:
            replacements["VERDICT"] = "N/A"
            replacements["VERDICT_REASON"] = ""
    else:
        replacements.setdefault("VERDICT", "N/A")
        replacements.setdefault("VERDICT_REASON", "")

    # Replace text placeholders
    replace_text_placeholders(doc, replacements)

    # Replace plot placeholders by embedding images (from index.json)
    replace_plot_placeholders(doc, index_entries, session_folder, max_width_in=5.5)

    # Replace artifact placeholders with md5 info
    art_pat = re.compile(r"\{\{ARTIFACT:([^}]+)\}\}")
    for p in doc.paragraphs:
        if not p.text:
            continue
        m = art_pat.search(p.text)
        if not m:
            continue
        name = m.group(1).strip()
        path = find_artifact_by_filename(index_entries, name)
        if path and path.exists():
            md5 = md5_file(path)
            p.text = p.text.replace(f"{{{{ARTIFACT:{name}}}}}", f"{path.name} (md5: {md5})")
        else:
            p.text = p.text.replace(f"{{{{ARTIFACT:{name}}}}}", f"{name} (missing)")

    # Replace table placeholders with DataFrames
    # Common tokens: MODEL_CV_RESULTS, PERFORMANCE_METRICS, FAIRNESS_METRICS, BOOTSTRAP_CI_SUMMARY, ARTIFACTS_LIST, LEAKAGE_TOP_FEATURES
    # Map tokens to data
    table_map = {
        "MODEL_CV_RESULTS": None,
        "PERFORMANCE_METRICS": results_df[[c for c in results_df.columns if c.lower() in ["model"] + [x.lower() for x in results_df.columns]]].copy() if not results_df.empty else results_df,
        "FAIRNESS_METRICS": results_df[[c for c in results_df.columns if c != "Model"]].copy() if not results_df.empty else results_df,
        "MODEL_PRIMARY_METRICS": results_df[[c for c in results_df.columns if c in ["Model","Accuracy","Disparate Impact","Statistical Parity Difference"]]] if not results_df.empty else results_df,
        "BOOTSTRAP_CI_SUMMARY": None,
        "ARTIFACTS_LIST": pd.DataFrame([{"file": e.get("file"), "kind": e.get("kind"), "path": Path(e.get("abs_path")).name if e.get("abs_path") else ""} for e in index_entries]) if index_entries else pd.DataFrame(),
        "LEAKAGE_TOP_FEATURES": None,
        "DECISION_THRESHOLDS": None,
    }

    # Insert tables for tokens present
    all_table_tokens = set([m.group(1) for p in doc.paragraphs for m in PLACEHOLDER_RE.finditer(p.text) if m.group(1).startswith("TABLE:")])
    # but simpler: check for each key in table_map and insert where placeholder exists
    for key, df in table_map.items():
        token = f"TABLE:{key}"
        if any(token in p.text for p in doc.paragraphs):
            insert_table_at_placeholder(doc, key, df)

    # Save output docx
    ts = datetime.now().strftime("%Y%m%dT%H%M%S")
    out_name = f"{output_prefix}_{ts}.docx"
    out_path = session_folder / out_name
    doc.save(str(out_path))

    # Write report metadata
    try:
        meta = {
            "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "file": out_name,
            "path": str(out_path),
            "md5": md5_file(out_path),
            "session": session_folder.name,
            "template": str(template_path.name)
        }
        (session_folder / "report_metadata.json").write_text(json.dumps(meta, indent=2), encoding="utf-8")
    except Exception:
        pass

    # Append to captures index.json
    try:
        idx_path = session_folder / "index.json"
        entries = []
        if idx_path.exists():
            entries = json.loads(idx_path.read_text(encoding="utf-8"))
        entries.append({"page":"Report","title":"AI Fairness Report","file":out_name,"abs_path":str(out_path.resolve()),"timestamp":ts,"kind":"docx"})
        idx_path.write_text(json.dumps(entries, indent=2), encoding="utf-8")
    except Exception:
        pass

    # Optionally log to MLflow
    if _MLFLOW_AVAILABLE:
        try:
            mlflow.set_experiment("Fairness_Eval")
        except Exception:
            pass
        try:
            run_name = f"report_{session_folder.name}_{ts}"
            with mlflow.start_run(run_name=run_name):
                # log some params from inf_meta
                for k, v in inf_meta.items():
                    try:
                        mlflow.log_param(str(k), str(v))
                    except Exception:
                        pass
                mlflow.set_tag("report_generated", "true")
                mlflow.log_artifact(str(out_path), artifact_path="reports")
        except Exception as e:
            print("MLflow logging failed:", e)

    print("Saved report:", out_path)
    return out_path

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--session", help="session id folder under captures (e.g., 123abc)", default=None)
    parser.add_argument("--folder", help="full path to captures/<session>", default=None)
    parser.add_argument("--template", help="custom template filename", default="AI_Fairness_Report_Template.docx")
    args = parser.parse_args()

    if args.folder:
        session_folder = Path(args.folder)
    elif args.session:
        session_folder = Path("captures") / args.session
    else:
        raise SystemExit("Provide --session <id> or --folder /path/to/captures/<session>")

    if not session_folder.exists():
        raise SystemExit(f"Session folder not found: {session_folder}")

    out = main(session_folder, template_name=args.template)
    print("Done. Report at:", out)
